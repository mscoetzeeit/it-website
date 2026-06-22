#!/usr/bin/env python
# Dump a question paper or memo (docx or pdf) to plain text for term mining.
import sys, os

def dump_docx(path):
    import docx
    d = docx.Document(path)
    out = []
    # interleave paragraphs and tables roughly in document order
    for p in d.paragraphs:
        t = p.text.strip()
        if t:
            out.append(t)
    for ti, tbl in enumerate(d.tables):
        for r in tbl.rows:
            cells = [c.text.strip().replace('\n', ' ') for c in r.cells]
            line = ' | '.join(c for c in cells if c)
            if line.strip():
                out.append(line)
    return '\n'.join(out)

def dump_pdf(path):
    from pypdf import PdfReader
    r = PdfReader(path)
    return '\n'.join((p.extract_text() or '') for p in r.pages)

if __name__ == '__main__':
    path = sys.argv[1]
    txt = dump_docx(path) if path.lower().endswith('.docx') else dump_pdf(path)
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    print(txt)
